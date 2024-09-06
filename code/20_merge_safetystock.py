import pandas as pd
from pathlib import Path
import xlsxwriter
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
# let's create a set of locals referring to our directory and working directory 
home_dir = Path.home()
work_dir = (home_dir / 'safety_stocks')
data = (work_dir / 'data')
raw_data = (data / 'raw')
code = Path.cwd() 
output = (work_dir / 'output')

# define function for loading CSV files into pandas dataframes 
def load_data():
    return {
        'bbg_rack_retail': pd.read_csv(f'{data}/bbg_rack_retail.csv'),
        'colonial_tariff': pd.read_csv(f'{data}/colonial_pipeline_tariffs_wide.csv'),
        'cpi': pd.read_csv(f'{data}/cpi.csv'),
        'eia_refiner_diesel': pd.read_csv(f'{data}/eia_refiner_diesel_prices.csv'),
        'eia_refiner_gasoline': pd.read_csv(f'{data}/eia_refiner_gasoline_prices.csv'),
        'eia_retail': pd.read_csv(f'{data}/eia_retail_prices.csv'),
        'eia_spot': pd.read_csv(f'{data}/eia_spot_prices.csv'),
        'eia_weekly_stock': pd.read_csv(f'{data}/eia_weekly_stock.csv'),
        'eia_monthly_stock': pd.read_csv(f'{data}/eia_monthly_stock.csv'),
        'eia_supplier_sales': pd.read_csv(f'{data}/eia_supplier_sales.csv')
    }

# define a cpi merge function 
def cpi_merge(df, cpi_df):
    df = pd.merge(df, cpi_df, on='date', how='outer')
    cpi_anchor = pd.to_datetime('2023-03-01')
    fixed_cpi = df.loc[df['date'] == cpi_anchor, 'all-urban cpi'].values
    df['price deflator'] = df['all-urban cpi'] / fixed_cpi
    if 'price deflator' not in df.columns:
        raise KeyError("'price deflator' column is missing after merge")
    return df 

def process_data(original_df, name_prefix, cpi_df):
    # Convert date columns to datetime
    original_df['date'] = pd.to_datetime(original_df['date'])
    original_df['month_year'] = original_df['date'].dt.to_period('M')  # Create month-year column
    
    # Collapse on month-year and calculate the mean for each group
    original_df = original_df.groupby('month_year').mean().reset_index()
    
    # Restore the 'date' column as the first of the month (since it was grouped by month-year)
    original_df['date'] = original_df['month_year'].dt.to_timestamp()
    original_df = original_df.drop('month_year', axis=1)

    # Apply CPI merge
    detailed_df = cpi_merge(original_df, cpi_df)
    
    # Forward fill the tariff columns directly in detailed_df
    cols_ffill = [col for col in detailed_df.columns if 'tariff' in col]
    
    if cols_ffill:
        # Forward-fill only the tariff columns in detailed_df
        detailed_df[cols_ffill] = detailed_df[cols_ffill].ffill()
    
    detailed_df_name = f'{name_prefix}_detailed_df'
    
    # Copy detailed_df back into original_df and drop CPI-related columns
    original_df = detailed_df.copy()
    original_df.drop(['price deflator', 'all-urban cpi'], axis=1, inplace=True)

    return original_df, detailed_df, detailed_df_name

def calculate_real_prices(df, nominal_cols):
    if 'price deflator' not in df.columns:
        raise KeyError("'price deflator' column is missing. Ensure CPI merge was successful.")
    for col in nominal_cols:
        real_col = col.replace(' (nominal)', ' (real)')
        df[real_col] = df[col] / df['price deflator']
    return df

def save_to_excel(df_dict, file_name):
    with pd.ExcelWriter(file_name, engine='xlsxwriter') as writer:
        for df_name, df in df_dict.items():
            sheet_name = df_name.replace('_detailed_df', '')[:31]
            df.to_excel(writer, sheet_name=sheet_name, index=False)

def process_master_data(dfs, cpi_df):
    master_df = pd.DataFrame()
    detailed_dfs = {}  # Dictionary to hold detailed DataFrames

    for df_name, df in dfs.items():
        processed_df, detailed_df, detailed_df_name = process_data(df, df_name, cpi_df)
        nominal_cols = [col for col in processed_df.columns if 'nominal' in col]
        
        detailed_df = calculate_real_prices(detailed_df, nominal_cols) if nominal_cols else detailed_df
        
        detailed_dfs[detailed_df_name] = detailed_df  # Store detailed_df in the dictionary

        if master_df.empty:
            master_df = processed_df.copy()
        else:
            master_df = pd.merge(master_df, processed_df, on='date', how='outer')

    return master_df, detailed_dfs

def main():
    dfs = load_data()
    cpi_df = dfs.pop('cpi')  # Remove CPI as it will be merged separately
    cpi_df['date'] = pd.to_datetime(cpi_df['date'])
    master_df, detailed_dfs = process_master_data(dfs, cpi_df)
    master_df = cpi_merge(master_df, cpi_df)

    save_to_excel(detailed_dfs, f'{data}/safety_stocks_master_detailed.xlsx')
    
    master_df['usgc/ny average low sulfur no. 2 spot (nominal)'] = (
        master_df['ny harbor-no. 2 diesel low sulfur spot price (nominal)'] 
        + master_df['usgc-no. 2 diesel low sulfur spot price (nominal)'])/2
    master_df['padd1a diesel retail spread (nominal)'] = (
        master_df['padd1a-no. 2 diesel low sulfur retail price (nominal)'] 
        - master_df['usgc/ny average low sulfur no. 2 spot (nominal)'])

    master_filtered_df = master_df[['date', 'padd1a diesel retail spread (nominal)', 
                                    'padd1a-no. 2 diesel low sulfur retail price (nominal)', 
                                    'usgc-no. 2 diesel low sulfur spot price (nominal)',
                                    'padd1a-distillate-weekly ending stock', 
                                    'houston-linden-buckeye tariff rate (nominal)', 
                                    'price deflator']]
    
    master_filtered_df.to_csv(f'{data}/safety_stocks_master.csv', index=False)
    return master_filtered_df, detailed_dfs

master_filtered_df, detailed_dfs = main()

def insert_graph_to_word(doc, image_path, plot_title):
    # Add title and image to the document
    doc.add_heading(plot_title, level=1)
    doc.add_paragraph(f"Graph: {plot_title}")
    doc.add_picture(image_path, width=Inches(5))

# define a plotting function
def plot_time_series(df, vars, title, ylabel, save_path):
    df['date'] = pd.to_datetime(df['date'])
    df = df.sort_values(by='date')
    plt.figure(figsize=(10, 6))
    for var in vars:
        # Plot the raw data
        plt.plot(df['date'], df[var], linestyle='-', 
                 linewidth=1, label=var)
        # Plot the 4-month moving average
        df[f'{var}_4m_MA'] = df[var].rolling(window=4, min_periods=1).mean()
        '''plt.plot(df['date'], df[f'{var}_4m_MA'], linestyle='-', 
                 linewidth=1.5, label=f'{var} (4-Month MA)')'''
    plt.legend(loc='upper left')
    plt.title(title)
    plt.xlabel('Date')
    plt.ylabel(ylabel)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(f'{output}/{save_path}.png')
    #plt.show()
    plt.close()
    insert_graph_to_word(doc, f'{output}/{save_path}.png', title)

doc = Document()
doc.add_heading('Data Inventory Time-Series', 0)

sheet_names = ['rack_retail prices, bbg', 'colonial pipeline rates',
               'refiner diesel prices, eia', 'refiner gasoline prices, eia',
               'retail prices, eia', 'spot prices, eia',
               'weekly stock, eia', 'monthly stock, eia',
               'supplier sales, eia']
graph_titles = ['Prices, BBG', 'Colonial Pipeline Tariff Rates',
                'Refiner Diesel Prices, EIA', 'Refiner Gasoline Prices, EIA',
                'Retail Prices, EIA', 'Spot Prices, EIA',
                'Weekly Ending Stock, EIA', 'Monthly Ending Stock, EIA', 
                'Supplier Sales, EIA']
graph_save_paths = ['rack_retail_prices_bbg', 'colonial_pipeline_tariff_rates',
                    'refiner_diesel_prices_eia', 'refiner_gasoline_prices_eia',
                    'retail_prices_eia', 'spot_prices_eia',
                    'weekly_stock_eia', 'stock_eia', 'supplier_sales_eia']

# now plot!
for df_name, df_plot, df_title, plot_title, plot_savepath in zip(
    detailed_dfs.keys(), detailed_dfs.values(), sheet_names, graph_titles, graph_save_paths):
    '''print(df_name) 
    print(df_plot.columns)'''
    # Replace these with string comparisons to match the keys in detailed_dfs
    if df_name in ['colonial_tariff_detailed_df', 'eia_spot_detailed_df', 
                   'bbg_rack_retail_detailed_df', 'eia_refiner_diesel_detailed_df', 
                   'eia_refiner_gasoline_detailed_df', 'eia_retail_detailed_df']:
        # Find the nominal columns for real price calculation
        vars_plot = [col.replace(' (nominal)', '') for col in df_plot.columns if 'nominal' in col]
        for var in vars_plot:
            nominal_var = f'{var} (nominal)'
            if nominal_var in df_plot.columns:
                df_plot[f'{var} (real)'] = (df_plot[nominal_var].dropna() 
                                          / df_plot['price deflator'].dropna())
        vars_plot_real = [var for var in df_plot.columns if '(real)' in var]
        if vars_plot_real:
            if df_name in ['colonial_tariff_detailed_df', 'eia_spot_detailed_df']:
                plot_time_series(df_plot, vars_plot_real, plot_title, 
                                 '$/Gal', plot_savepath)
            if df_name == 'bbg_rack_retail_detailed_df':
                dyed_columns = [col for col in vars_plot_real if 'heating oil dyed' in col]
                if dyed_columns:
                    plot_time_series(df_plot, dyed_columns, f'Heating Oil, Dyed Rack {plot_title}',
                                     '$/Gal', f'heating_dyed_{plot_savepath}')
                residential_columns = [col for col in vars_plot_real
                                       if 'heating oil residential price' in col]
                if residential_columns:
                    plot_time_series(df_plot, residential_columns, 
                                     f'Residential Heating Oil Retail {plot_title}', 
                                     '$/Gal', f'heating_resident_{plot_savepath}')
            if df_name == 'eia_refiner_diesel_detailed_df':
                retail_columns = [col for col in vars_plot_real if 'retail' in col]
                if retail_columns:
                    plot_time_series(df_plot, retail_columns, f'Retail {plot_title}',
                                     '$/Gal', f'retail_{plot_savepath}')
                wholesale_resale_columns = [col for col in vars_plot_real if 'wholesale/resale' in col]
                if wholesale_resale_columns:
                    plot_time_series(df_plot, wholesale_resale_columns, 
                                     f'Wholesale/Resale {plot_title}', '$/Gal', 
                                     f'wholesale_resale_{plot_savepath}')
            if df_name == 'eia_refiner_gasoline_detailed_df':
                company_outlets_columns = [col for col in vars_plot_real 
                                           if 'company outlets' in col and 'total' in col]
                if company_outlets_columns:
                    plot_time_series(df_plot, company_outlets_columns, 
                                     f'Through Company Outlets {plot_title}', 
                                     '$/Gal', f'company_outlets_{plot_savepath}')
                other_users = [col for col in vars_plot_real if 'other users' in col and 'total' in col]
                if other_users:
                    plot_time_series(df_plot, other_users, f'Other Users {plot_title}', 
                                     '$/Gal', f'other_users_{plot_savepath}')
                wholesale_resale_columns = [col for col in vars_plot_real if 'wholesale/resale' in col]
                if wholesale_resale_columns:
                    plot_time_series(df_plot, wholesale_resale_columns, 
                                     f'Wholesale/Resale {plot_title}', 
                                     '$/Gal', f'wholesale_resale_{plot_savepath}')
    if df_name == 'eia_monthly_stock_detailed_df':
        vars_plot = [col for col in df_plot.columns if 'stock' in col]
        distillate_column = [col for col in vars_plot if 'distillate' in col]
        if distillate_column:
            plot_time_series(df_plot, distillate_column, f'Distillate {plot_title}',
                         'Barrels', f'distillate_{plot_savepath}')
        conventional_mg_column = [col for col in vars_plot if 'conventional mg' in col]
        if conventional_mg_column:
            plot_time_series(df_plot, conventional_mg_column, 
                         f'Conventional Motor Gasoline {plot_title}',
                        'Barrels', f'conventional_mg_{plot_savepath}')
        reformulated_mg_column = [col for col in vars_plot if 'rfmg' in col]
        if reformulated_mg_column:
            plot_time_series(df_plot, reformulated_mg_column, 
                         f'Reformulated Motor Gasoline {plot_title}', 'Barrels', 
                         f'rfmg_{plot_savepath}')
    if df_name == 'eia_weekly_stock_detailed_df': 
        vars_plot = [col for col in df_plot.columns if 'stock' in col]
        plot_time_series(df_plot, vars_plot, plot_title, 'Barrels', plot_savepath)
    if df_name == 'eia_supplier_sales_detailed_df':
        vars_plot = [col for col in df_plot.columns if 'sales' in col]
        distillate_column = [col for col in vars_plot if 'distillate' in col]
        if distillate_column:
            plot_time_series(df_plot, distillate_column, f'Distillate {plot_title}', 
                         'Barrels', f'distillate_{plot_savepath}')
        fuel_heating_column = [col for col in vars_plot if 'fuel/heating oil' in col]
        if fuel_heating_column:
            plot_time_series(df_plot, fuel_heating_column, 
                         f'Fuel/Heating Oil {plot_title}', 'Barrels', 
                         f'fuel_heating_{plot_savepath}')
        diesel_column = [col for col in vars_plot if ('diesel' in col) and ('low sulfur' not in col)]
        if diesel_column:
            plot_time_series(df_plot, diesel_column, 
                         f'Diesel {plot_title}', 'Barrels', 
                         f'diesel_{plot_savepath}')
        diesel_low_sulfur_column = [col for col in vars_plot if 'diesel' in col and 'low sulfur' in col]
        if diesel_low_sulfur_column:
            plot_time_series(df_plot, diesel_low_sulfur_column, 
                         f'Diesel Low Sulfur {plot_title}', 'Barrels', 
                         f'diesel_lowsulfur_{plot_savepath}')
doc.save(f'{output}/data_inventory_compiled_timeseries.docx')
print("Document saved as 'data_inventory_compiled_timeseries.docx'")